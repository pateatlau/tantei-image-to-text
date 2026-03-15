"""
Convert Markdown to DOCX
=========================
Converts the Markdown OCR output to a Word document with proper formatting.

Usage:
    python convert_to_docx.py

Input: output_proofread.md (or output.md if proofread doesn't exist)
Output: output.docx

Dependencies:
    pip install python-docx markdown
"""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import markdown
from html.parser import HTMLParser

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------
INPUT_PATH = Path("output_proofread.md")
FALLBACK_PATH = Path("output.md")
OUTPUT_PATH = Path("output.docx")


# ---------------------------------------------------------------------------
# Markdown to DOCX converter
# ---------------------------------------------------------------------------
class MarkdownToDocx:
    def __init__(self):
        self.doc = Document()
        self.setup_styles()

    def setup_styles(self):
        """Set default document margins and styles."""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    def add_heading_1(self, text):
        """Add a level 1 heading (# in Markdown)."""
        heading = self.doc.add_heading(text, level=1)
        heading.runs[0].font.size = Pt(16)
        heading.runs[0].font.bold = True

    def add_heading_2(self, text):
        """Add a level 2 heading (## in Markdown)."""
        heading = self.doc.add_heading(text, level=2)
        heading.runs[0].font.size = Pt(14)
        heading.runs[0].font.bold = True

    def add_heading_3(self, text):
        """Add a level 3 heading (### in Markdown)."""
        heading = self.doc.add_heading(text, level=3)
        heading.runs[0].font.size = Pt(12)
        heading.runs[0].font.bold = True

    def add_paragraph(self, text, indent=0):
        """Add a regular paragraph with optional indentation."""
        p = self.doc.add_paragraph(text)
        p.paragraph_format.left_indent = Inches(indent * 0.25)
        p.paragraph_format.space_after = Pt(6)

    def add_bullet(self, text, indent=0):
        """Add a bullet point with optional indentation."""
        p = self.doc.add_paragraph(text, style='List Bullet')
        p.paragraph_format.left_indent = Inches(indent * 0.25)

    def add_numbered(self, text, indent=0):
        """Add a numbered list item with optional indentation."""
        p = self.doc.add_paragraph(text, style='List Number')
        p.paragraph_format.left_indent = Inches(indent * 0.25)

    def parse_markdown_line(self, line, prev_line_type=None):
        """Parse a single line of Markdown and add to document."""
        stripped = line.strip()

        if not stripped:
            return None

        # Detect headings
        if stripped.startswith('# '):
            self.add_heading_1(stripped[2:])
            return 'heading1'
        elif stripped.startswith('## '):
            self.add_heading_2(stripped[3:])
            return 'heading2'
        elif stripped.startswith('### '):
            self.add_heading_3(stripped[4:])
            return 'heading3'

        # Detect bullet points
        elif re.match(r'^[\-\*]\s+', stripped):
            text = re.sub(r'^[\-\*]\s+', '', stripped)
            indent = len(line) - len(line.lstrip())
            self.add_bullet(text, indent // 2)
            return 'bullet'

        # Detect numbered lists
        elif re.match(r'^\d+\.\s+', stripped):
            text = re.sub(r'^\d+\.\s+', '', stripped)
            indent = len(line) - len(line.lstrip())
            self.add_numbered(text, indent // 2)
            return 'numbered'

        # Regular paragraph
        else:
            indent = len(line) - len(line.lstrip())
            self.add_paragraph(stripped, indent // 2)
            return 'paragraph'

    def convert(self, markdown_text):
        """Convert full Markdown text to DOCX."""
        lines = markdown_text.split('\n')
        prev_type = None

        for line in lines:
            prev_type = self.parse_markdown_line(line, prev_type)

    def save(self, output_path):
        """Save the document."""
        self.doc.save(output_path)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main():
    # Determine input file
    if INPUT_PATH.exists():
        input_file = INPUT_PATH
        print(f"Using proofread version: {input_file}")
    elif FALLBACK_PATH.exists():
        input_file = FALLBACK_PATH
        print(f"Proofread version not found, using: {input_file}")
    else:
        print(f"Error: No input file found. Run ocr_book.py first.")
        sys.exit(1)

    # Read Markdown
    print(f"Reading Markdown from {input_file}...")
    markdown_text = input_file.read_text(encoding='utf-8')

    # Convert to DOCX
    print(f"Converting to DOCX...")
    converter = MarkdownToDocx()
    converter.convert(markdown_text)
    converter.save(OUTPUT_PATH)

    print(f"\nDone! Word document saved to {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
