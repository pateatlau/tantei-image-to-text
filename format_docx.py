"""
Format DOCX — Add table borders, header, and footer
=====================================================
Post-processes a pandoc-generated DOCX to add:
1. Visible borders on all tables
2. Header with document title + thin bottom border
3. Footer with "Page X of Y pages" + thin top border

Usage:
    python format_docx.py [input.docx]

Defaults:
    Input/Output: output_final.docx (modified in place)

Dependencies:
    pip install python-docx
"""

import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

DEFAULT_PATH = Path("output_final.docx")
TITLE = "ETHICS IN ADMINISTRATION + PROBITY + GOVERNANCE"
BORDER_COLOR = "999999"


def set_table_borders(table):
    """Add single-line borders to all sides of a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')

    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), BORDER_COLOR)
        borders.append(element)

    # Remove existing borders if any
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)

    tblPr.append(borders)


def add_paragraph_border(paragraph, edge='bottom'):
    """Add a thin border to a paragraph (top or bottom)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    border = OxmlElement(f'w:{edge}')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), '4')
    border.set(qn('w:space'), '8')
    border.set(qn('w:color'), BORDER_COLOR)
    pBdr.append(border)
    pPr.append(pBdr)


def add_field(paragraph, field_code):
    """Add a Word field (e.g. PAGE, NUMPAGES) to a paragraph."""
    run = paragraph.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar_begin)

    run2 = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f' {field_code} '
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run3._r.append(fldChar_end)


def setup_header(doc):
    """Add title with bottom border to the document header."""
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False

    # Clear existing header content
    for p in header.paragraphs:
        p.clear()

    # Use first paragraph or add one
    if header.paragraphs:
        para = header.paragraphs[0]
    else:
        para = header.add_paragraph()

    run = para.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add spacing after header paragraph
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:after'), '200')
    pPr.append(spacing)

    add_paragraph_border(para, 'bottom')


def setup_footer(doc):
    """Add 'Page X of Y pages' with top border to the document footer."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    # Clear existing footer content
    for p in footer.paragraphs:
        p.clear()

    if footer.paragraphs:
        para = footer.paragraphs[0]
    else:
        para = footer.add_paragraph()

    # Add spacing before footer paragraph
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '200')
    pPr.append(spacing)

    add_paragraph_border(para, 'top')

    # "Page "
    run = para.add_run("Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # PAGE field
    add_field(para, 'PAGE')

    # " of "
    run = para.add_run(" of ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # NUMPAGES field
    add_field(para, 'NUMPAGES')

    # " pages"
    run = para.add_run(" pages")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main():
    docx_path = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_PATH

    if not docx_path.exists():
        print(f"Error: {docx_path} not found.")
        sys.exit(1)

    print(f"Formatting {docx_path}...")

    doc = Document(str(docx_path))

    # 1. Table borders
    table_count = len(doc.tables)
    for table in doc.tables:
        set_table_borders(table)
    print(f"  ✓ Added borders to {table_count} tables")

    # 2. Header
    setup_header(doc)
    print(f"  ✓ Added header: \"{TITLE}\"")

    # 3. Footer
    setup_footer(doc)
    print(f"  ✓ Added footer: \"Page X of Y pages\"")

    doc.save(str(docx_path))
    print(f"\n✓ Saved {docx_path}")


if __name__ == "__main__":
    main()
